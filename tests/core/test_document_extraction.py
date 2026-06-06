# tests/core/test_document_extraction.py
"""Bateria de testes para src/core/document_extraction.py.

Cobre extração de texto pra ingestionar KB. Cada formato (docx, pdf, md)
produz mesmo output normalizado quando entrada é equivalente.

Fixtures dinâmicas (gera arquivos em runtime no tmp_path) — sem assets
binários commitados.
"""
from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest


# ───────────────────────────────────────────────────────────────────────
# _clean (normalizador)
# ───────────────────────────────────────────────────────────────────────


class TestClean:
    """Normalização: CRLF → LF, colapsa newlines, colapsa whitespace, strip."""

    @pytest.mark.parametrize("raw,expected", [
        ("line1\r\nline2", "line1\nline2"),       # CRLF → LF
        ("line1\rline2", "line1\nline2"),         # CR → LF
        ("aaa\n\n\n\nbbb", "aaa\n\nbbb"),         # colapsa 3+ newlines
        ("aaa    bbb", "aaa bbb"),                 # colapsa spaces
        ("aaa\t\tbbb", "aaa bbb"),                 # colapsa tabs
        ("   leading", "leading"),                 # strip leading
        ("trailing   ", "trailing"),               # strip trailing
        ("", ""),
    ])
    def test_normalization(self, raw, expected):
        from src.core.document_extraction import _clean
        assert _clean(raw) == expected


# ───────────────────────────────────────────────────────────────────────
# extract_md
# ───────────────────────────────────────────────────────────────────────


class TestExtractMd:

    def test_from_filesystem_path(self, tmp_path):
        from src.core.document_extraction import extract_md
        p = tmp_path / "doc.md"
        p.write_text("# Hello\n\nCompêndio test", encoding="utf-8")
        text = extract_md(p)
        assert "Hello" in text
        assert "Compêndio" in text

    def test_from_bytesio_upload(self):
        from src.core.document_extraction import extract_md
        content = b"# Title\n\nSome content"
        stream = BytesIO(content)
        text = extract_md(stream)
        assert "Title" in text
        assert "Some content" in text

    def test_utf8_decoded_correctly(self, tmp_path):
        from src.core.document_extraction import extract_md
        p = tmp_path / "doc.md"
        p.write_bytes("Açaí 🌙 cabelo crespo".encode("utf-8"))
        text = extract_md(p)
        assert "Açaí" in text
        assert "🌙" in text

    def test_fallback_latin1_when_not_utf8(self, tmp_path):
        """Arquivo latin-1 vintage não quebra a ingestão."""
        from src.core.document_extraction import extract_md
        p = tmp_path / "doc.md"
        # Latin-1 puro (não-UTF-8 válido)
        p.write_bytes(b"ol\xe1 mundo")
        text = extract_md(p)
        assert "ol" in text

    def test_applies_clean_to_md_content(self, tmp_path):
        """_clean é aplicado mesmo em markdown (whitespace normalizado)."""
        from src.core.document_extraction import extract_md
        p = tmp_path / "doc.md"
        p.write_text("para1\n\n\n\n\n\npara2  com  espaços  duplos", encoding="utf-8")
        text = extract_md(p)
        assert "\n\n\n" not in text
        assert "  " not in text


# ───────────────────────────────────────────────────────────────────────
# extract_docx (precisa criar docx temporário)
# ───────────────────────────────────────────────────────────────────────


def _make_docx(path: Path, paragraphs: list[str], tables: list[list[list[str]]] | None = None):
    """Cria docx mínimo pra testar extração."""
    from docx import Document
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if tables:
        for table_data in tables:
            t = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            for r, row in enumerate(table_data):
                for c, cell in enumerate(row):
                    t.cell(r, c).text = cell
    doc.save(str(path))


class TestExtractDocx:

    def test_extracts_paragraphs(self, tmp_path):
        from src.core.document_extraction import extract_docx
        p = tmp_path / "test.docx"
        _make_docx(p, ["Primeira linha", "Segunda linha", "Terceira"])
        text = extract_docx(str(p))
        assert "Primeira linha" in text
        assert "Segunda linha" in text
        assert "Terceira" in text

    def test_extracts_tables_with_pipe_separator(self, tmp_path):
        from src.core.document_extraction import extract_docx
        p = tmp_path / "test.docx"
        _make_docx(
            p,
            ["Antes da tabela"],
            tables=[[["Coluna A", "Coluna B"], ["valor 1", "valor 2"]]],
        )
        text = extract_docx(str(p))
        assert "Antes da tabela" in text
        # Tabela formatada como "cell1 | cell2"
        assert "Coluna A | Coluna B" in text
        assert "valor 1 | valor 2" in text

    def test_skips_empty_paragraphs(self, tmp_path):
        from src.core.document_extraction import extract_docx
        p = tmp_path / "test.docx"
        _make_docx(p, ["", "  ", "Real content", ""])
        text = extract_docx(str(p))
        assert text == "Real content"  # vazios e whitespace puro foram strippados


# ───────────────────────────────────────────────────────────────────────
# extract_pdf
# ───────────────────────────────────────────────────────────────────────


def _make_pdf(path: Path, text: str) -> None:
    """Cria PDF mínimo via pypdf não cobre criação — usamos um workaround
    com reportlab se disponível, senão skip.
    """
    try:
        from reportlab.pdfgen import canvas  # type: ignore[import]
    except ImportError:
        pytest.skip("reportlab not installed, skipping PDF generation")

    c = canvas.Canvas(str(path))
    y = 800
    for line in text.split("\n"):
        c.drawString(72, y, line)
        y -= 20
    c.save()


class TestExtractPdf:

    def test_extracts_simple_text(self, tmp_path):
        from src.core.document_extraction import extract_pdf
        p = tmp_path / "test.pdf"
        _make_pdf(p, "Página 1 linha A\nPágina 1 linha B")
        text = extract_pdf(str(p))
        # pypdf extrai texto, ordem pode variar — só checa presença
        assert "Página 1 linha A" in text or "linha A" in text


# ───────────────────────────────────────────────────────────────────────
# extract_text (router por extensão)
# ───────────────────────────────────────────────────────────────────────


class TestExtractText:

    def test_routes_md_by_extension(self, tmp_path):
        from src.core.document_extraction import extract_text
        p = tmp_path / "compendio.md"
        p.write_text("Markdown content", encoding="utf-8")
        text = extract_text("compendio.md", path=p)
        assert "Markdown content" in text

    def test_routes_markdown_alias(self, tmp_path):
        from src.core.document_extraction import extract_text
        p = tmp_path / "compendio.markdown"
        p.write_text("Also markdown", encoding="utf-8")
        text = extract_text("compendio.markdown", path=p)
        assert "Also markdown" in text

    def test_routes_docx_by_extension(self, tmp_path):
        from src.core.document_extraction import extract_text
        p = tmp_path / "doc.docx"
        _make_docx(p, ["Docx routed correctly"])
        text = extract_text("doc.docx", path=p)
        assert "Docx routed correctly" in text

    def test_accepts_inmemory_bytes_for_upload(self, tmp_path):
        """Upload via API passa `data` em bytes, não path no disco."""
        from src.core.document_extraction import extract_text
        md_bytes = b"# Upload test\n\nUploaded content"
        text = extract_text("upload.md", data=md_bytes)
        assert "Upload test" in text
        assert "Uploaded content" in text

    def test_unsupported_extension_raises(self):
        from src.core.document_extraction import extract_text
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_text("file.txt", data=b"hello")

    @pytest.mark.parametrize("filename", [
        "DOC.MD",       # uppercase extension
        "Doc.Md",       # mixed case
        "doc.MD",
    ])
    def test_case_insensitive_routing(self, filename, tmp_path):
        from src.core.document_extraction import extract_text
        text = extract_text(filename, data=b"# Hello\n")
        assert "Hello" in text
