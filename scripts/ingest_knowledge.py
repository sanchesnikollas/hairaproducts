#!/usr/bin/env python3
"""Ingest the Doutoras knowledge base (docx/pdf) into normalized .txt files.

Reads `data/knowledge_base/*.{docx,pdf}`, extracts clean text and saves to
`data/knowledge_base/processed/<slug>.txt`. Moon loads these at startup.

This is intentionally "inject everything" (~36k tokens total) — small enough to
fit comfortably in Claude's 200k context with Anthropic prompt caching. RAG
becomes worth it only when the corpus exceeds ~100k tokens.

Usage:
    python scripts/ingest_knowledge.py            # process all files
    python scripts/ingest_knowledge.py --list     # list what would be processed
"""
from __future__ import annotations

import os
import re
import sys
import unicodedata
from pathlib import Path

from docx import Document
from pypdf import PdfReader

KB_DIR = Path("data/knowledge_base")
OUT_DIR = KB_DIR / "processed"

# Files to skip (e.g., superseded versions). Keys are basenames.
SKIP = {
    "Haira - regras 2.pdf",  # superseded by Haira-Regras-3.0.docx
}


def _slug(name: str) -> str:
    n = unicodedata.normalize("NFD", name)
    n = "".join(c for c in n if unicodedata.category(c) != "Mn")
    n = re.sub(r"[^a-zA-Z0-9]+", "_", n).strip("_").lower()
    return n


def _extract_docx(path: Path) -> str:
    d = Document(str(path))
    parts: list[str] = []
    for p in d.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_pdf(path: Path) -> str:
    r = PdfReader(str(path))
    pages = []
    for p in r.pages:
        t = p.extract_text() or ""
        t = re.sub(r"[ \t]+", " ", t)
        pages.append(t.strip())
    return "\n\n".join(pages)


def _clean(text: str) -> str:
    # Collapse extra whitespace but keep paragraph breaks
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def main() -> None:
    if "--list" in sys.argv:
        for fn in sorted(os.listdir(KB_DIR)):
            if fn in SKIP:
                print(f"  SKIP  {fn}")
            elif fn.endswith((".docx", ".pdf")):
                print(f"  PICK  {fn}")
        return

    # Extract + upsert into the local knowledge_chunks table. The processed
    # /processed/*.txt files are kept as a side-effect for debugging/diffing
    # but NEVER committed (data/knowledge_base is gitignored).
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    from sqlalchemy.orm import Session
    from src.storage.database import get_engine
    from src.storage.knowledge_models import KnowledgeChunkORM

    rows: list[tuple[str, str, int]] = []
    for fn in sorted(os.listdir(KB_DIR)):
        if fn in SKIP:
            print(f"SKIP (superseded): {fn}")
            continue
        path = KB_DIR / fn
        if path.is_dir():
            continue
        if fn.endswith(".docx"):
            text = _extract_docx(path)
        elif fn.endswith(".pdf"):
            text = _extract_pdf(path)
        else:
            continue
        text = _clean(text)
        out = OUT_DIR / f"{_slug(path.stem)}.txt"
        out.write_text(text, encoding="utf-8")
        rows.append((fn, text, len(text)))
        print(f"OK  {fn:<35} -> {out.name} ({len(text)} chars, ~{int(len(text.split())*1.3)} tokens)")

    # Upsert into DB (the source of truth Moon reads from)
    engine = get_engine()
    with Session(engine) as s:
        existing = {r.source for r in s.query(KnowledgeChunkORM.source).all()}
        for source, content, n in rows:
            if source in existing:
                row = s.get(KnowledgeChunkORM, source)
                row.content = content
                row.char_count = n
            else:
                s.add(KnowledgeChunkORM(source=source, content=content, char_count=n))
        # remove anything no longer in the source folder
        present = {src for src, _, _ in rows}
        for stale in existing - present:
            s.delete(s.get(KnowledgeChunkORM, stale))
        s.commit()
    total_chars = sum(n for _, _, n in rows)
    print(f"\nTOTAL: {len(rows)} fontes, {total_chars} chars, ~{int(total_chars/4)} tokens (estim.)")
    print("Persisted into knowledge_chunks table. Sync to prod via scripts/sync_knowledge_base.py.")


if __name__ == "__main__":
    main()
