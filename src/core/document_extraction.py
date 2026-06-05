"""Document text extraction for the knowledge base (docx, pdf, md).

Used by `scripts/ingest_knowledge.py` (file-based ingestion) and by the admin
upload endpoint (`POST /api/admin/knowledge/upload`). Centralised so both paths
produce the same normalized text.

Markdown was added in 2026-06-04 to ingest the Compêndio Haira (entregue como
.md exportado do Google Docs). Não converte para HTML — preserva o texto
como está e deixa que Anthropic interprete os marcadores naturalmente.
"""
from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


def _clean(text: str) -> str:
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def extract_docx(stream_or_path) -> str:
    d = Document(stream_or_path)
    parts: list[str] = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _clean("\n".join(parts))


def extract_pdf(stream_or_path) -> str:
    r = PdfReader(stream_or_path)
    pages = []
    for p in r.pages:
        t = p.extract_text() or ""
        t = re.sub(r"[ \t]+", " ", t)
        pages.append(t.strip())
    return _clean("\n\n".join(pages))


def extract_md(stream_or_path) -> str:
    """Read markdown verbatim. Accepts BytesIO (upload) or a filesystem path."""
    if hasattr(stream_or_path, "read"):
        raw = stream_or_path.read()
    else:
        raw = Path(stream_or_path).read_bytes()
    if isinstance(raw, bytes):
        # Compêndio em UTF-8; fallback latin-1 só pra não quebrar.
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            text = raw.decode("latin-1", errors="replace")
    else:
        text = raw
    # Markdown já vem semi-limpo; aplica o normalizador padrão.
    return _clean(text)


def extract_text(filename: str, data: bytes | None = None, path: Path | None = None) -> str:
    """Extract from a filename (for routing by extension), accepting either
    in-memory bytes (upload) or a filesystem path (script)."""
    fn = filename.lower()
    src = BytesIO(data) if data is not None else str(path)
    if fn.endswith(".docx"):
        return extract_docx(src)
    if fn.endswith(".pdf"):
        return extract_pdf(src)
    if fn.endswith(".md") or fn.endswith(".markdown"):
        return extract_md(src)
    raise ValueError(f"Unsupported file type: {filename} (only .docx, .pdf, .md)")
